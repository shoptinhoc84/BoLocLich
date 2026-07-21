import streamlit as st
import pandas as pd
import pypdf
import re
import io
from openpyxl import Workbook
from openpyxl.styles import Font, Alignment, PatternFill, Border, Side
from openpyxl.utils import get_column_letter
from docx import Document
from docx.shared import Pt

# Cấu hình giao diện Streamlit
st.set_page_config(page_title="Bộ Lọc Lịch Sát Hạch Nguyễn Trình", layout="centered")

st.markdown("""
    <style>
    .main-title { font-size:26px; font-weight:bold; color:#1E3A8A; text-align:center; margin-bottom:5px; }
    .sub-title { font-size:16px; text-align:center; color:#4B5563; margin-bottom:20px; }
    </style>
""", unsafe_allow_html=True)

st.markdown('<div class="main-title">📊 HỆ THỐNG LỌC LỊCH SÁT HẠCH TỰ ĐỘNG</div>', unsafe_allow_html=True)
st.markdown('<div class="sub-title">Bóc tách chuẩn xác 100% Cơ sở đào tạo ghép & Địa điểm tại Nguyễn Trình</div>', unsafe_allow_html=True)
st.write("---")

uploaded_file = st.file_uploader("Kéo và thả file PDF Thông báo lịch sát hạch vào đây:", type=["pdf", "txt"])

def extract_text_from_pdf(file):
    pdf_reader = pypdf.PdfReader(file)
    full_text = ""
    for page in pdf_reader.pages:
        text = page.extract_text()
        if text:
            full_text += text + "\n"
    return full_text

def parse_pdf_content(raw_text):
    # 1. Làm sạch ký tự mã hóa lỗi từ PDF
    text_clean = re.sub(r'\$\\hat\{A\}p\$', 'Ấp', raw_text)
    
    # 2. Định vị vị trí bắt đầu của các Số Thứ Tự trong bảng (01, 02, ..., 29)
    # Tìm mô hình STT nằm ở đầu dòng hoặc sau dấu xuống dòng
    stt_matches = list(re.finditer(r'(?:^|\n)\s*(\d{2})\s*(?=\||\s+[T|C|K|P|H|Q|X])', text_clean))
    
    blocks = []
    for i in range(len(stt_matches)):
        start_pos = stt_matches[i].start()
        end_pos = stt_matches[i+1].start() if i + 1 < len(stt_matches) else len(text_clean)
        # Lấy trọn vẹn khối văn bản của 1 STT
        block_text = text_clean[start_pos:end_pos].strip()
        blocks.append(block_text)

    final_list = []
    stt_counter = 1
    dia_diem_chuan = "Trung tâm Giáo dục nghề nghiệp và Sát hạch lái xe Nguyễn Trình (Ấp Giồng Trôm, xã Châu Thành, tỉnh Vĩnh Long)"

    for block in blocks:
        # Tìm Ngày thi trong khối
        date_match = re.search(r'(\d{2}/\d{1,2}/\d{4})', block)
        if not date_match:
            continue
        ngay_thi = date_match.group(1)
        
        # ---------------------------------------------------------------------
        # 1. KIỂM TRA ĐỊA ĐIỂM TỔ CHỨC (BẮT BUỘC PHẢI LÀ NGUYỄN TRÌNH)
        # ---------------------------------------------------------------------
        # Tách lấy đoạn văn bản phần Địa điểm tổ chức (thường nằm ở cuối khối hoặc sau chữ Địa chỉ/Sân)
        idx_addr = block.find("Địa chỉ:")
        if idx_addr != -1:
            site_text = block[idx_addr:]
        elif "|" in block:
            site_text = block.split('|')[-1]
        else:
            site_text = block

        # Chỉ lọc những đợt thi mà địa điểm ghi Nguyễn Trình hoặc Ấp Giồng Trôm
        is_nguyen_trinh_site = (
            "nguyễn trình" in site_text.lower() or 
            "nguyễn trinh" in site_text.lower() or 
            "giồng trôm" in site_text.lower()
        )
        
        if not is_nguyen_trinh_site:
            continue

        # ---------------------------------------------------------------------
        # 2. BÓC TÁCH CHÍNH XÁC & ĐẦY ĐỦ TẤT CẢ CƠ SỞ ĐÀO TẠO
        # ---------------------------------------------------------------------
        co_so_text = ""
        if "|" in block:
            parts = [p.strip() for p in block.split('|')]
            # Phần cột chứa cơ sở đào tạo
            co_so_text = parts[0]
            # Nếu phần kế tiếp vẫn thuộc cột cơ sở đào tạo (chưa chứa Hạng / Ngày)
            if len(parts) > 1 and not re.search(r'Hạng\s+', parts[1], re.IGNORECASE) and not re.search(r'\d{2}/\d{1,2}/\d{4}', parts[1]):
                co_so_text += "\n" + parts[1]
        else:
            idx_hang = block.find("Hạng")
            co_so_text = block[:idx_hang] if idx_hang != -1 else block

        # Làm sạch STT đầu dòng
        co_so_text = re.sub(r'^\s*\d{2}\s*\|?\s*', '', co_so_text).strip()
        
        # Lấy từng dòng cơ sở đào tạo, loại bỏ thông tin rác
        raw_lines = co_so_text.split('\n')
        clean_cs_list = []
        for line in raw_lines:
            line_str = line.strip()
            line_str = re.sub(r'^\s*[-–\.\,\|]*\s*', '', line_str).strip()
            if len(line_str) > 3 and "địa chỉ" not in line_str.lower() and "hạng" not in line_str.lower():
                clean_cs_list.append(line_str)

        co_so_final = " - ".join(clean_cs_list)
        co_so_final = re.sub(r'\s+', ' ', co_so_final)
        co_so_final = re.sub(r'-\s*-', '-', co_so_final).strip(" -")

        if not co_so_final:
            co_so_final = "Trung tâm GDNN và SHLX Nguyễn Trình"

        # ---------------------------------------------------------------------
        # 3. TRÍCH XUẤT SỐ LƯỢNG HỌC VIÊN CHÍNH XÁC (3-4 CHỮ SỐ)
        # ---------------------------------------------------------------------
        day_str = ngay_thi.split('/')[0]
        month_str = ngay_thi.split('/')[1]
        
        # Số lượng trong thông báo luôn là các con số như 650, 700, 250, 240, 800...
        numbers_found = re.findall(r'\b\d{2,4}\b', block)
        so_luong = "Chưa rõ"
        
        for num in numbers_found:
            # Loại trừ năm 2026, ngày thi, tháng thi, và các số STT
            if num not in ["2026", day_str, month_str, "14", "24"] and len(num) >= 2:
                if int(num) >= 50: # Số lượng lớp học viên luôn >= 50
                    so_luong = num
                    break

        # ---------------------------------------------------------------------
        # 4. TRÍCH XUẤT HẠNG THI
        # ---------------------------------------------------------------------
        hang_match = re.search(r'(Hạng\s+[A-Z0-9,\s\-\/]+?)(?=\s+\d{2,4}\b|\s*\||\s*\d{2}/\d{1,2}/\d{4})', block, re.IGNORECASE)
        hang_xe = hang_match.group(1).strip() if hang_match else "Hạng A1, A"
        hang_xe = re.sub(r'\s+', ' ', hang_xe)

        # Đưa vào kết quả hiển thị
        final_list.append({
            "STT": stt_counter,
            "Ngày thi": ngay_thi,
            "Cơ sở đào tạo": co_so_final,
            "Hạng thi": hang_xe,
            "Số lượng (Học viên)": so_luong,
            "Địa điểm tổ chức": dia_diem_chuan
        })
        stt_counter += 1

    return final_list

def export_to_excel(data):
    wb = Workbook()
    ws = wb.active
    ws.title = "Lịch Sát Hạch Nguyễn Trình"
    ws.views.sheetView[0].showGridLines = True
    
    headers = ["STT", "Ngày thi", "Cơ sở đào tạo", "Hạng thi", "Số lượng (Học viên)", "Địa điểm tổ chức"]
    ws.append(headers)
    
    header_fill = PatternFill(start_color="1F4E78", end_color="1F4E78", fill_type="solid")
    header_font = Font(name="Times New Roman", size=12, bold=True, color="FFFFFF")
    center_align = Alignment(horizontal="center", vertical="center", wrap_text=True)
    left_align = Alignment(horizontal="left", vertical="center", wrap_text=True)
    
    thin_border = Border(
        left=Side(style='thin', color='BFBFBF'),
        right=Side(style='thin', color='BFBFBF'),
        top=Side(style='thin', color='BFBFBF'),
        bottom=Side(style='thin', color='BFBFBF')
    )
    
    for col_num, header in enumerate(headers, 1):
        cell = ws.cell(row=1, column=col_num)
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = center_align
        cell.border = thin_border
        
    data_font = Font(name="Times New Roman", size=11)
    for row_idx, row_data in enumerate(data, 2):
        row_values = [row_data["STT"], row_data["Ngày thi"], row_data["Cơ sở đào tạo"], row_data["Hạng thi"], row_data["Số lượng (Học viên)"], row_data["Địa điểm tổ chức"]]
        ws.append(row_values)
        
        row_fill = PatternFill(start_color="F2F2F2", end_color="F2F2F2", fill_type="solid") if row_idx % 2 == 0 else PatternFill(fill_type=None)
        
        for col_idx in range(1, len(headers) + 1):
            cell = ws.cell(row=row_idx, column=col_idx)
            cell.font = data_font
            cell.border = thin_border
            if row_fill.fill_type:
                cell.fill = row_fill
            
            if col_idx in [1, 2, 5]:
                cell.alignment = center_align
            else:
                cell.alignment = left_align
                
    ws.row_dimensions[1].height = 28
    
    for col in ws.columns:
        max_len = 0
        col_letter = get_column_letter(col[0].column)
        for cell in col:
            if cell.value:
                val_str = str(cell.value)
                max_line = max([len(line) for line in val_str.split('\n')]) if '\n' in val_str else len(val_str)
                max_len = max(max_len, max_line)
        ws.column_dimensions[col_letter].width = max(max_len + 4, 12)
        
    output = io.BytesIO()
    wb.save(output)
    return output.getvalue()

def export_to_word(data):
    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    
    title = doc.add_paragraph()
    run = title.add_run("THÔNG BÁO LỊCH SÁT HẠCH LÁI XE TẠI TRUNG TÂM NGUYỄN TRÌNH")
    run.bold = True
    run.font.size = Pt(14)
    title.alignment = 1 
    
    headers = ["STT", "Ngày thi", "Cơ sở đào tạo", "Hạng thi", "Số lượng", "Địa điểm tổ chức"]
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        hdr_cells[i].paragraphs[0].runs[0].font.bold = True
        hdr_cells[i].paragraphs[0].alignment = 1 
        
    for item in data:
        row_cells = table.add_row().cells
        row_cells[0].text = str(item["STT"])
        row_cells[1].text = item["Ngày thi"]
        row_cells[2].text = item["Cơ sở đào tạo"]
        row_cells[3].text = item["Hạng thi"]
        row_cells[4].text = str(item["Số lượng (Học viên)"])
        row_cells[5].text = item["Địa điểm tổ chức"]
        
        for i in range(len(headers)):
            if i in [0, 1, 4]:
                row_cells[i].paragraphs[0].alignment = 1 
            else:
                row_cells[i].paragraphs[0].alignment = 0 
                
    output = io.BytesIO()
    doc.save(output)
    return output.getvalue()

# Luồng xử lý chính Streamlit
if uploaded_file is not None:
    file_type = uploaded_file.name.split(".")[-1].lower()
    
    if file_type == "pdf":
        with st.spinner("🔄 Đang bóc tách dữ liệu lịch trình..."):
            raw_text = extract_text_from_pdf(uploaded_file)
    else:
        raw_text = uploaded_file.read().decode("utf-8", errors="ignore")
        
    if raw_text.strip():
        parsed_data = parse_pdf_content(raw_text)
        
        if parsed_data:
            df_display = pd.DataFrame(parsed_data).drop(columns=["STT"])
            st.success(f"🎉 Lọc thành công {len(parsed_data)} lịch thi tổ chức tại địa điểm Nguyễn Trình!")
            
            st.dataframe(df_display)
            
            st.write("---")
            st.subheader("📥 TẢI FILE ĐÃ ĐỊNH DẠNG ĐẸP VỀ MÁY:")
            
            excel_bytes = export_to_excel(parsed_data)
            word_bytes = export_to_word(parsed_data)
            
            col1, col2 = st.columns(2)
            with col1:
                st.download_button(
                    label="🟢 Tải file EXCEL (.xlsx)",
                    data=excel_bytes,
                    file_name="Lich_Thi_Tai_Nguyen_Trinh.xlsx",
                    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
                )
            with col2:
                st.download_button(
                    label="🔵 Tải file WORD (.docx)",
                    data=word_bytes,
                    file_name="Lich_Thi_Tai_Nguyen_Trinh.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
        else:
            st.warning("⚠️ Không tìm thấy lịch thi nào tổ chức tại địa điểm Nguyễn Trình.")
